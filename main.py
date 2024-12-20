from datetime import datetime
import os
from docx import Document
from dotenv import load_dotenv, dotenv_values
import telebot 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.styles.style import BaseStyle

from consts import BASE_TEMPLATE_FOLDER

import os.path

import database
from telegram_templates import document_keyboard

load_dotenv()
config = dotenv_values(".env")

API_TOKEN = config['TOKEN']

bot = telebot.TeleBot(API_TOKEN)

# Handle '/start' and '/help'
@bot.message_handler(commands=['help', 'start'])
def send_welcome(message):
    bot.reply_to(message, """\
/glipceb ФИО - Лист исполнения на почту, 1С, ЕОСДО, BOXER
/glip ФИО - Лист исполнения на почту
/glic ФИО - Лист исполнения на 1С
/glice ФИО - Лист исполнения на 1С, ЕОСДО
/glipc ФИО - Лист исполнения на 1С, Почту
/glie ФИО - Лист исполнения на ЕОСДО
/gliep ФИО - Лист исполнения на ЕОСДО, Почту
/glib ФИО - Лист исполнения на BOXER
/gsz ФИО - Служебка на УЗ в kisozk.local
/glipce ФИО - Лист исполнения на почту, 1С, ЕОСДО
/glicb ФИО - Лист исполнения на 1С, BOXER
/gliceb ФИО - Лист исполнения на 1С, ЕОСДО, BOXER
/glipcb ФИО - Лист исполнения на 1С, Почту, BOXER
/glieb ФИО - Лист исполнения на ЕОСДО, BOXER
/gliepb ФИО - Лист исполнения на ЕОСДО, Почту, BOXER
/glics ФИО - Лист исполнения на 1С, Станция сканирования
/glies ФИО - Лист исполнения на ЕОСДО, Станция сканирования
/glips ФИО - Лист исполнения на почту, Станция сканирования
/gli ФИО - Лист исполнения на почту, 1С, ЕОСДО, BOXER, Станция сканирования
/glices ФИО - Лист исполнения на 1С, ЕОСДО, Станция сканирования
/glipcs ФИО - Лист исполнения на 1С, Почту, Станция сканирования
/glibs ФИО - Лист исполнения на BOXER, Станция сканирования
/glieps ФИО - Лист исполнения на ЕОСДО, Почту, Станция сканирования
/glipces ФИО - Лист исполнения на почту, 1С, ЕОСДО, Станция сканирования
/glicbs ФИО - Лист исполнения на 1С, BOXER, Станция сканирования
/glicebs ФИО - Лист исполнения на 1С, ЕОСДО, BOXER, Станция сканирования
/glipcbs ФИО - Лист исполнения на 1С, Почту, BOXER, Станция сканирования
/gliebs ФИО - Лист исполнения на ЕОСДО, BOXER, Станция сканирования
/gliepbs ФИО - Лист исполнения на ЕОСДО, Почту, BOXER, Станция сканирования
""")

def apply_style(paragraph: Paragraph, text: str, style: BaseStyle):
    paragraph.text = text
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_styles(doc):
    styles = doc.styles
    BIGstyle = styles.add_style('Big', WD_STYLE_TYPE.PARAGRAPH)
    BIGstyle.font.name = 'Times New Roman'
    BIGstyle.font.size = Pt(12)

    return BIGstyle


#list ispolneniya na 1C
@bot.callback_query_handler(func=lambda call: call.data == 'glic')
def get_worker(call):
    if not database.check_permissions(call.from_user.id):
        bot.reply_to(call, "Доступа нет")
        return
    print(call.from_user.id)
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(call, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_1C.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C.doc"
    doc.save(filename)
    bot.send_document(call.message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO
@bot.callback_query_handler(func=lambda call: call.data == 'glie')
def get_worker(call):
    if not database.check_permissions(call.from_user.id):
        bot.reply_to(call, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(call, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_EOSDO.doc"
    doc.save(filename)
    bot.send_document(call.message.chat.id, open(filename, 'rb'))
    os.remove(filename)


#list ispolneniya na vse krome SS
@bot.callback_query_handler(func=lambda call: call.data == 'glipceb')
def get_worker(call):
    if not database.check_permissions(call.from_user.id):
        bot.reply_to(call, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(call, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+1C+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_Pochta_Eosdo_1C_BOXER.doc"
    doc.save(filename)
    bot.send_document(call.message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#sluzhebka na kisozk
@bot.message_handler(commands=['gsz'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "SZ_istok_na_UZ.docx"))
    BIGstyle = build_styles(doc)
    
    apply_style(doc.tables[1].rows[1].cells[0].paragraphs[0], fio, BIGstyle)
    apply_style(doc.tables[1].rows[1].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[1].rows[1].cells[2].paragraphs[0], str(number), BIGstyle)

    apply_style(doc.tables[2].rows[0].cells[0].paragraphs[0] , position, BIGstyle)
    apply_style(doc.tables[2].rows[0].cells[1].paragraphs[0] , fio, BIGstyle)

    filename = f"SZ_{name}_kisozk.local.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)


#list ispolneniya na pochtu
@bot.message_handler(commands=['glip'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_Pochta.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)


#list ispolneniya na 1C EOSDO
@bot.message_handler(commands=['glice'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+1C.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_EOSDO.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C+Pochta
@bot.message_handler(commands=['glipc'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+1C.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_Pochta.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO+Pochta
@bot.message_handler(commands=['gliep'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_EOSDO_Pochta.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na BOXER
@bot.message_handler(commands=['glib'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Boxer.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[0], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_BOXER.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C eosdo boxer
@bot.message_handler(commands=['gliceb'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+1C+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}_Eosdo_1C_BOXER.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C boxer
@bot.message_handler(commands=['glicb'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_1C+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_BOXER.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

# list ispolneniya na pochta 1c boxer
@bot.message_handler(commands=['glipcb'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+1C+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}1C_BOXER_Pochta.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO+BOxer
@bot.message_handler(commands=['glieb'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_EOSDO_BOXER.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO Pochta BOXER
@bot.message_handler(commands=['gliepb'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+BOXER.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}_Eosdo_Pochta_BOXER.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C eosdo pochtu
@bot.message_handler(commands=['glipce'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+1C.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)


    filename = f"LI_{name}_Eosdo_1C_pochta.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na vse
@bot.message_handler(commands=['gli'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+1C+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[10].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[10].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[10].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[10].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[10].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[10].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_Pochta_Eosdo_1C_BOXER_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C SS
@bot.message_handler(commands=['glics'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_1C+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO SS
@bot.message_handler(commands=['glies'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[3].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[3].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[3].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[3].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[3].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[3].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_EOSDO_SS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na pochtu ss
@bot.message_handler(commands=['glips'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[3].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[3].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[3].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[3].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[3].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[3].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_Pochta_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C EOSDO ss
@bot.message_handler(commands=['glices'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+1C+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_EOSDO_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C+Pochta ss
@bot.message_handler(commands=['glipcs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+1C+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_Pochta_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na BOXER ss
@bot.message_handler(commands=['glibs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Boxer+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[0], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_BOXER_SS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO+Pochta ss
@bot.message_handler(commands=['glieps'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)
        
    filename = f"LI_{name}_EOSDO_Pochta_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C eosdo pochtu ss
@bot.message_handler(commands=['glipces'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+1C+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)


    filename = f"LI_{name}_Eosdo_1C_pochta_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C boxer ss
@bot.message_handler(commands=['glicbs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_1C+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_1C_BOXER_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na 1C eosdo boxer ss
@bot.message_handler(commands=['glicebs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+1C+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}_Eosdo_1C_BOXER_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

# list ispolneniya na pochta 1c boxer ss
@bot.message_handler(commands=['glipcbs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+1C+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}1C_BOXER_Pochta_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO+BOxer ss
@bot.message_handler(commands=['gliebs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_EOSDO+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)
    
    filename = f"LI_{name}_EOSDO_BOXER_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

#list ispolneniya na EOSDO Pochta BOXER ss
@bot.message_handler(commands=['gliepbs'])
def get_worker(message):
    if not database.check_permissions(message.from_user.id):
        bot.reply_to(message, "Доступа нет")
        return
    fio_from_user=database.get_fio_from_user(call.from_user.id)[0]
    data = database.select_from_datauser(fio_from_user)
    if not data:
        bot.reply_to(message, "Работник не найден")
        return
    data = data[0] # TODO: поставить обработку нескольких работников

    fio = data[1]
    number = data[2]
    position = data[3]
    department = data[4]
    address = data[5]
    director = data[6]
    date = datetime.today().strftime("%d.%m.%Y")
    name = fio.split(" ") 
    name = name[0] + " " + name[1][0] + "." + (name[2][0] + "." if len(name) > 2 else "")

    doc = Document(os.path.join(BASE_TEMPLATE_FOLDER, "LI_Pochta+EOSDO+BOXER+SS.docx"))
    BIGstyle = build_styles(doc)

    apply_style(doc.tables[0].rows[1].cells[1].paragraphs[1], fio, BIGstyle)
    apply_style(doc.tables[0].rows[5].cells[1].paragraphs[0], str(number), BIGstyle)
    apply_style(doc.tables[0].rows[6].cells[1].paragraphs[1], position, BIGstyle)
    apply_style(doc.tables[0].rows[9].cells[1].paragraphs[0], department, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[1], address, BIGstyle)
    apply_style(doc.tables[0].rows[10].cells[1].paragraphs[0], '', BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[1].paragraphs[1], name, BIGstyle)
    apply_style(doc.tables[0].rows[15].cells[4].paragraphs[1], date, BIGstyle)
    
    apply_style(doc.tables[2].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[2].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[2].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[2].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[4].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[4].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[4].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[4].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[6].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[6].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[6].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[6].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    apply_style(doc.tables[8].rows[8].cells[2].paragraphs[0], director, BIGstyle)
    apply_style(doc.tables[8].rows[8].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[9].cells[4].paragraphs [0], date, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[1].paragraphs[0], name, BIGstyle)
    apply_style(doc.tables[8].rows[17].cells[3].paragraphs[0], date, BIGstyle)
    apply_style(doc.tables[8].rows[3].cells[4].paragraphs [0], date, BIGstyle)

    filename = f"LI_{name}_Eosdo_Pochta_BOXER_StS.doc"
    doc.save(filename)
    bot.send_document(message.chat.id, open(filename, 'rb'))
    os.remove(filename)

@bot.message_handler()
def priem_fio(message):
    if database.check_users_by_fio(message.text):
        database.add_fio(message.text, message.from_user.id)
    else:
        bot.send_message(message.chat.id, "работник не найден")
        return
    bot.send_message(message.chat.id, 'Выберите ИТ-ресурсы', reply_markup=document_keyboard())


bot.infinity_polling()
